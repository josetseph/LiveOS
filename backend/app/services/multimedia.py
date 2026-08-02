"""Multimedia processing via local model services and lightweight document parsers."""

# pylint: disable=wrong-import-order,import-outside-toplevel
import os
import csv
from collections.abc import Callable

from app.core.config import settings
from app.core.log import get_logger

logger = get_logger("MultimediaService")


def _format_timestamp(seconds: float) -> str:
    """Convert seconds to M:SS or H:MM:SS string."""
    s = int(seconds)
    m, s = divmod(s, 60)
    h, m = divmod(m, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m}:{s:02d}"


class MultimediaService:
    """Extract text from attachments using local model services."""

    def _parse_storage_ref(self, path_or_url: str) -> tuple[str, str] | None:
        """Return (bucket, key) when the path points at RustFS / proxy storage."""
        from urllib.parse import urlparse

        if path_or_url.startswith("http"):
            parsed = urlparse(path_or_url)
            parts = parsed.path.strip("/").split("/", 1)
            if len(parts) == 2 and parts[1]:
                return parts[0], parts[1]
            return None

        if path_or_url.startswith("/files/"):
            remainder = path_or_url[len("/files/") :].lstrip("/")
            parts = remainder.split("/", 1)
            if len(parts) == 2 and parts[1]:
                return parts[0], parts[1]
            return None

        if path_or_url.startswith("/uploads/"):
            key = path_or_url[len("/uploads/") :].lstrip("/")
            if key:
                return settings.BUCKET_NAME, key

        files_url = settings.FILES_URL.rstrip("/")
        # Desktop uses FILES_URL=/vault-files as a local vault proxy — never S3.
        if (
            getattr(settings, "STORAGE_BACKEND", "local") == "s3"
            and files_url.startswith("/")
            and path_or_url.startswith(files_url + "/")
            and not path_or_url.startswith("/vault-files/")
        ):
            key = path_or_url[len(files_url) + 1 :]
            if key:
                return settings.BUCKET_NAME, key

        return None

    def _download_from_storage(self, bucket: str, key: str) -> str:
        """Download an object from RustFS/S3 to a temporary local file."""
        import tempfile

        import boto3
        from botocore.client import Config

        suffix = "." + key.rsplit(".", 1)[-1] if "." in key else ".tmp"
        client = boto3.client(
            "s3",
            aws_access_key_id=settings.BUCKET_ACCESS_KEY_ID,
            aws_secret_access_key=settings.BUCKET_SECRET_ACCESS_KEY,
            endpoint_url=settings.R2_ENDPOINT_URL,
            config=Config(s3={"addressing_style": "path"}),
        )
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            client.download_fileobj(bucket, key, tmp)
            return tmp.name

    def _resolve_storage_url(self, path_or_url: str) -> str:
        """Normalize attachment URLs/paths to something the backend can fetch."""
        if not path_or_url:
            return path_or_url

        if path_or_url.startswith("http"):
            return path_or_url

        if os.path.isfile(path_or_url):
            return path_or_url

        if path_or_url.startswith("/files/"):
            remainder = path_or_url[len("/files/") :].lstrip("/")
            if remainder:
                return f"{settings.R2_ENDPOINT_URL.rstrip('/')}/{remainder}"

        if path_or_url.startswith("/uploads/"):
            key = path_or_url[len("/uploads/") :].lstrip("/")
            if key:
                return (
                    f"{settings.R2_ENDPOINT_URL.rstrip('/')}/"
                    f"{settings.BUCKET_NAME}/{key}"
                )

        files_url = settings.FILES_URL.rstrip("/")
        if files_url.startswith("/") and path_or_url.startswith(files_url + "/"):
            key = path_or_url[len(files_url) + 1 :]
            if key:
                return (
                    f"{settings.R2_ENDPOINT_URL.rstrip('/')}/"
                    f"{settings.BUCKET_NAME}/{key}"
                )

        return path_or_url

    def _resolve_vault_local_path(self, path_or_url: str) -> str | None:
        """Map ``/vault-files/<kb_id>/<rel>`` to an absolute vault file path."""
        from pathlib import Path
        from urllib.parse import unquote, urlparse

        if not path_or_url:
            return None

        raw = path_or_url.strip()
        if raw.startswith("http://") or raw.startswith("https://"):
            raw = urlparse(raw).path or ""

        raw = raw.split("?", 1)[0].replace("\\", "/")
        marker = "/vault-files/"
        if marker not in raw and not raw.startswith("vault-files/"):
            return None

        if raw.startswith("vault-files/"):
            raw = "/" + raw
        idx = raw.find(marker)
        remainder = raw[idx + len(marker) :] if idx >= 0 else ""
        parts = remainder.split("/", 1)
        if len(parts) < 2 or not parts[0] or not parts[1]:
            return None

        kb_id, rel = parts[0], unquote(parts[1]).lstrip("/")
        if not rel or ".." in Path(rel).parts:
            return None

        try:
            from app.services.kb_registry import kb_registry

            kb = kb_registry.get_kb(kb_id)
        except Exception as exc:  # pylint: disable=broad-exception-caught
            logger.warning(f"Vault lookup failed for kb={kb_id}: {exc}")
            return None

        if not kb or not kb.vault_path:
            logger.warning(f"No vault for kb={kb_id} (url={path_or_url})")
            return None

        vault_root = Path(kb.vault_path).resolve()
        abs_path = (vault_root / rel).resolve()
        try:
            abs_path.relative_to(vault_root)
        except ValueError:
            logger.warning(f"Rejected path escape: {abs_path} not under {vault_root}")
            return None

        if abs_path.is_file():
            return str(abs_path)
        logger.warning(f"Vault file missing: {abs_path}")
        return None

    def _is_ephemeral_download(self, original_ref: str, local_path: str) -> bool:
        """True when ``local_path`` is a temp download we must delete (not a vault file)."""
        if not local_path or local_path == original_ref:
            return False
        if self._resolve_vault_local_path(original_ref) == local_path:
            return False
        if os.path.isfile(original_ref) and os.path.samefile(original_ref, local_path):
            return False
        return True

    def _download_temp_file(self, path_or_url: str) -> str:
        """Resolve vault/local/storage/remote attachments to a local file path."""
        import tempfile

        import requests

        if os.path.isfile(path_or_url):
            return path_or_url

        # Vault attachments are local-only — never fall through to RustFS/S3.
        if "/vault-files/" in path_or_url or path_or_url.lstrip("/").startswith(
            "vault-files/"
        ):
            vault_path = self._resolve_vault_local_path(path_or_url)
            if vault_path:
                logger.info(f"Resolved vault attachment: {vault_path}")
                return vault_path
            raise FileNotFoundError(f"Vault attachment not found: {path_or_url}")

        vault_path = self._resolve_vault_local_path(path_or_url)
        if vault_path:
            logger.info(f"Resolved vault attachment: {vault_path}")
            return vault_path

        if getattr(settings, "STORAGE_BACKEND", "local") != "s3":
            if os.path.isfile(path_or_url):
                return path_or_url
            raise FileNotFoundError(f"Attachment not found: {path_or_url}")

        storage_ref = self._parse_storage_ref(path_or_url)
        if storage_ref:
            bucket, key = storage_ref
            logger.info(f"Downloading storage object: {bucket}/{key}")
            return self._download_from_storage(bucket, key)

        resolved = self._resolve_storage_url(path_or_url)
        if resolved != path_or_url:
            vault_path = self._resolve_vault_local_path(resolved)
            if vault_path:
                return vault_path
            storage_ref = self._parse_storage_ref(resolved)
            if storage_ref:
                bucket, key = storage_ref
                logger.info(f"Downloading storage object: {bucket}/{key}")
                return self._download_from_storage(bucket, key)

        if not resolved.startswith("http"):
            if os.path.isfile(resolved):
                return resolved
            raise FileNotFoundError(f"Attachment not found: {path_or_url}")

        logger.info(f"Downloading remote file: {resolved}...")
        response = requests.get(resolved, timeout=300)
        response.raise_for_status()
        suffix = "." + resolved.split(".")[-1] if "." in resolved else ".tmp"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(response.content)
            return tmp.name

    def _describe_image_local(self, local_path: str) -> str:
        """Florence caption via in-process multimodal runtime."""
        from app.services.multimodal_runtime import multimodal_runtime

        return multimodal_runtime.describe_image_path(local_path) or ""

    def _caption_video_with_marlin(self, local_path: str) -> dict:
        """Marlin caption via in-process multimodal runtime."""
        from app.services.multimodal_runtime import multimodal_runtime

        try:
            return multimodal_runtime.caption_video_path(local_path)
        except Exception as exc:  # pylint: disable=broad-exception-caught
            logger.error(f"Marlin captioning failed: {exc}")
            raise RuntimeError(f"Marlin captioning failed: {exc}") from exc

    def unload_local_models(self, family: str | None = None) -> None:
        """Unload Florence/Whisper (and optionally Marlin) from the API process."""
        from app.services.multimodal_runtime import multimodal_runtime

        try:
            # Legacy callers pass "florence" / "whisper"; marlin has its own unload.
            multimodal_runtime.unload(family)
        except Exception as exc:  # pylint: disable=broad-exception-caught
            logger.warning(f"Multimodal unload skipped/failed: {exc}")

    def unload_marlin(self) -> None:
        """Unload Marlin from the API process."""
        from app.services.multimodal_runtime import multimodal_runtime

        try:
            multimodal_runtime.unload("marlin")
        except Exception as exc:  # pylint: disable=broad-exception-caught
            logger.warning(f"Marlin unload skipped/failed: {exc}")

    def describe_image(self, image_path: str) -> str:
        """Generate an image description via Florence (local) or cloud vision fallback."""
        local_path = self._download_temp_file(image_path)
        try:
            try:
                text = self._describe_image_local(local_path)
                if text:
                    return text
            except Exception as exc:  # pylint: disable=broad-exception-caught
                logger.warning(f"Local image description failed: {exc}")
            # Cloud vision only when not in local-only AI mode.
            if (settings.AI_SETUP_MODE or "").lower() not in ("local", "none"):
                cloud = self._describe_image_cloud(local_path)
                if cloud:
                    return cloud
            raise RuntimeError("Image description failed (local Florence unavailable)")
        finally:
            if self._is_ephemeral_download(image_path, local_path) and os.path.exists(
                local_path
            ):
                os.remove(local_path)

    def _describe_image_cloud(self, image_path: str) -> str:
        """Optional OpenAI / Gemini vision when local Florence is unavailable."""
        try:
            import base64

            with open(image_path, "rb") as f:
                data = f.read()
            b64 = base64.b64encode(data).decode("ascii")
            mime = "image/png" if image_path.lower().endswith(".png") else "image/jpeg"
            data_url = f"data:{mime};base64,{b64}"

            if settings.OPENAI_API_KEY:
                from openai import OpenAI

                client = OpenAI(api_key=settings.OPENAI_API_KEY)
                resp = client.chat.completions.create(
                    model=settings.OPENAI_MODEL or "gpt-4o-mini",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "Describe this image briefly."},
                                {"type": "image_url", "image_url": {"url": data_url}},
                            ],
                        }
                    ],
                    max_tokens=400,
                )
                return (resp.choices[0].message.content or "").strip()

            if settings.GEMINI_API_KEY:
                from google import genai
                from google.genai import types

                client = genai.Client(api_key=settings.GEMINI_API_KEY)
                part = types.Part.from_bytes(data=data, mime_type=mime)
                resp = client.models.generate_content(
                    model=settings.GEMINI_MODEL or "gemini-2.0-flash",
                    contents=["Describe this image briefly.", part],
                )
                return (resp.text or "").strip()
        except Exception as exc:  # pylint: disable=broad-exception-caught
            logger.debug(f"Cloud vision failed: {exc}")
        return ""

    def transcribe_audio(self, audio_path: str) -> str:
        """Transcribe audio via the local-models service."""
        local_path = self._download_temp_file(audio_path)
        try:
            logger.info(f"Transcribing audio: {local_path}")
            from app.services.multimodal_runtime import multimodal_runtime

            return multimodal_runtime.transcribe_audio_path(local_path) or ""
        finally:
            if self._is_ephemeral_download(audio_path, local_path) and os.path.exists(
                local_path
            ):
                os.remove(local_path)

    def process_video(self, video_path: str) -> str:
        """Process a video with Whisper audio transcription and Marlin visual analysis."""
        transcript = self.transcribe_video_audio(video_path)
        visual = self.describe_video_visual(video_path)
        parts = []
        if transcript:
            parts.append(f"### Spoken Content\n{transcript}")
        if visual:
            parts.append(visual)
        return "\n\n".join(parts) if parts else "(Video processing produced no output)"

    def transcribe_video_audio(self, video_path: str) -> str:
        """Transcribe a video's audio track without running visual analysis."""
        import av

        local_path = self._download_temp_file(video_path)
        try:
            with av.open(local_path) as container:
                has_video = len(container.streams.video) > 0

            if not has_video:
                logger.info("No video stream detected — treating as audio-only.")
                return self.transcribe_audio(local_path)

            try:
                logger.info("Transcribing video audio track...")
                return self.transcribe_audio(local_path)
            except Exception as exc:  # pylint: disable=broad-exception-caught
                logger.warning(
                    f"Audio transcription skipped (no audio track or failed): {exc}"
                )
                return ""
        finally:
            if self._is_ephemeral_download(video_path, local_path) and os.path.exists(
                local_path
            ):
                os.remove(local_path)

    def describe_video_visual(self, video_path: str) -> str:
        """Run Marlin visual analysis without transcribing audio."""
        import av

        local_path = self._download_temp_file(video_path)
        try:
            with av.open(local_path) as container:
                has_video = len(container.streams.video) > 0

            if not has_video:
                return ""

            try:
                logger.info("Running Marlin video captioning via service...")
                result = self._caption_video_with_marlin(local_path)
                scene = result.get("scene", "")
                events = result.get("events", [])
                lines = [
                    f"- {_format_timestamp(ev.get('start', 0))}\u2013{_format_timestamp(ev.get('end', 0))} \u2014 {ev.get('description', '')}"
                    for ev in events
                ]
                events_text = "\n".join(lines)
                logger.info(f"Marlin: {len(events)} events extracted.")
            except Exception as exc:  # pylint: disable=broad-exception-caught
                logger.error(f"Marlin captioning failed: {exc}")
                return ""

            if scene:
                visual = f"### Visual Analysis\n**Scene:** {scene}"
                if events_text:
                    visual += f"\n\n**Events:**\n{events_text}"
                return visual
            return ""
        finally:
            if self._is_ephemeral_download(video_path, local_path) and os.path.exists(
                local_path
            ):
                os.remove(local_path)

    def _pdf_page_needs_render(self, page, native_text: str, image_descriptions: list[str]) -> bool:
        """True when a full-page Florence render is useful (scanned / sparse pages)."""
        if not settings.PDF_VISUAL_EXTRACTION_ENABLED:
            return False
        if len(native_text.strip()) >= settings.PDF_VISUAL_TEXT_THRESHOLD:
            return False
        # Embedded-image Florence already covered this page.
        if image_descriptions:
            return False
        try:
            if page.get_images(full=True):
                return True
        except Exception:  # pylint: disable=broad-exception-caught
            pass
        try:
            if page.get_drawings():
                return True
        except Exception:  # pylint: disable=broad-exception-caught
            pass
        # Image-only / empty pages with almost no text still benefit from a render.
        return len(native_text.strip()) == 0

    def _describe_pdf_page_render(self, page) -> str:
        """Render a PDF page to PNG and describe it with Florence."""
        import tempfile

        import fitz

        dpi = max(int(settings.PDF_VISUAL_RENDER_DPI or 144), 72)
        zoom = dpi / 72.0
        pixmap = page.get_pixmap(
            matrix=fitz.Matrix(zoom, zoom),
            alpha=False,
            annots=True,
        )
        image_path = ""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                image_path = tmp.name
                pixmap.save(image_path)
            result_text = self._describe_image_local(image_path)
            return (result_text or "").strip()
        finally:
            if image_path and os.path.exists(image_path):
                os.remove(image_path)

    def extract_text_from_pdf(
        self,
        pdf_path: str,
        progress_callback: Callable[[str, str | None], None] | None = None,
    ) -> str:
        """Extract PDF page text; Florence on embedded images and sparse page renders."""
        import tempfile

        import fitz

        def _progress(stage: str, model: str | None = None) -> None:
            if progress_callback:
                progress_callback(stage, model)

        local_path = self._download_temp_file(pdf_path)
        owns_temp = self._is_ephemeral_download(pdf_path, local_path)
        try:
            extracted_pages: list[str] = []
            doc = fitz.open(local_path)
            total_pages = len(doc)
            max_visual_pages = int(settings.PDF_VISUAL_EXTRACTION_MAX_PAGES or 0)
            visual_pages_used = 0
            try:
                for page_index, page in enumerate(doc, start=1):
                    _progress(
                        f"PDF: page {page_index}/{total_pages}, extracting text",
                        None,
                    )
                    native_text = page.get_text().strip()
                    page_parts = [f"--- Page {page_index} ---"]
                    image_descriptions: list[str] = []

                    if native_text:
                        page_parts.append(f"Native text:\n{native_text}")

                    images = page.get_images(full=True)
                    for image_index, image_info in enumerate(images, start=1):
                        _progress(
                            (
                                f"PDF: page {page_index}/{total_pages}, "
                                f"describing image {image_index}/{len(images)}"
                            ),
                            "Florence-2",
                        )
                        xref = image_info[0]
                        image_path = ""
                        try:
                            extracted = doc.extract_image(xref)
                            image_bytes = extracted.get("image")
                            if not image_bytes:
                                continue
                            image_ext = extracted.get("ext") or "png"
                            with tempfile.NamedTemporaryFile(
                                delete=False, suffix=f".{image_ext}"
                            ) as tmp:
                                image_path = tmp.name
                                tmp.write(image_bytes)
                            description = self._describe_image_local(image_path)
                        except Exception as exc:  # pylint: disable=broad-exception-caught
                            logger.warning(
                                "PDF image description skipped "
                                f"(page={page_index}, image={image_index}): {exc}"
                            )
                            continue
                        finally:
                            if image_path and os.path.exists(image_path):
                                os.remove(image_path)

                        if description:
                            image_descriptions.append(
                                f"Image {image_index}: {description}"
                            )

                    # Scanned / sparse pages: render whole page when embeds yielded nothing.
                    if self._pdf_page_needs_render(page, native_text, image_descriptions):
                        if not max_visual_pages or visual_pages_used < max_visual_pages:
                            _progress(
                                f"PDF: page {page_index}/{total_pages}, describing page render",
                                "Florence-2",
                            )
                            try:
                                page_desc = self._describe_pdf_page_render(page)
                                if page_desc:
                                    image_descriptions.append(
                                        f"Page render: {page_desc}"
                                    )
                                    visual_pages_used += 1
                            except Exception as exc:  # pylint: disable=broad-exception-caught
                                logger.warning(
                                    f"PDF page render description skipped "
                                    f"(page={page_index}): {exc}"
                                )

                    if image_descriptions:
                        page_parts.append(
                            "Image descriptions:\n" + "\n".join(image_descriptions)
                        )

                    if native_text or image_descriptions:
                        extracted_pages.append("\n\n".join(page_parts))

                    _progress(f"PDF: page {page_index}/{total_pages} complete", None)
            finally:
                doc.close()

            full_text = "\n\n".join(extracted_pages).strip()
            if not full_text:
                return "PDF contains no extractable native or visual content."
            return full_text
        except Exception as exc:  # pylint: disable=broad-exception-caught
            logger.error(f"PDF extraction failed: {exc}")
            raise RuntimeError(f"PDF extraction failed: {exc}") from exc
        finally:
            if owns_temp and local_path and os.path.exists(local_path):
                os.remove(local_path)

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from a Word document (.docx) using native parsing only."""
        local_path = self._download_temp_file(docx_path)
        parts = []

        try:
            try:
                import docx

                document = docx.Document(local_path)
                for paragraph in document.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        parts.append(text)

                for idx, table in enumerate(document.tables, start=1):
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        parts.append(f"--- Table {idx} ---\n" + "\n".join(rows))

                full_text = "\n\n".join(parts).strip()
                if not full_text:
                    return "Word document contains no extractable text."
                return full_text

            except ImportError as exc:
                raise RuntimeError(
                    "Word extraction unavailable: python-docx is not installed."
                ) from exc
            except Exception as exc:  # pylint: disable=broad-exception-caught
                logger.error(f"DOCX extraction failed: {exc}")
                raise RuntimeError(f"Word extraction failed: {exc}") from exc

        finally:
            if self._is_ephemeral_download(docx_path, local_path) and os.path.exists(
                local_path
            ):
                os.remove(local_path)

    def extract_text_from_spreadsheet(
        self, sheet_path: str
    ) -> str:  # pylint: disable=too-many-return-statements,too-many-nested-blocks,too-many-locals,too-many-branches
        """Extract text from spreadsheet-like files using native parsing."""
        local_path = self._download_temp_file(sheet_path)
        lower_path = local_path.lower()

        try:  # pylint: disable=too-many-nested-blocks
            if lower_path.endswith(".xlsx"):
                try:
                    from openpyxl import load_workbook

                    workbook = load_workbook(
                        filename=local_path,
                        read_only=True,
                        data_only=True,
                    )
                    parts = []
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        rows = []
                        for row in sheet.iter_rows(values_only=True):
                            values = [
                                str(cell).strip() for cell in row if cell is not None
                            ]
                            if values:
                                rows.append("\t".join(values))
                        if rows:
                            parts.append(
                                f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows)
                            )

                    workbook.close()
                    full_text = "\n\n".join(parts).strip()
                    if not full_text:
                        return "Spreadsheet contains no extractable text."
                    return full_text

                except ImportError as exc:
                    raise RuntimeError(
                        "Spreadsheet extraction unavailable: openpyxl is not installed."
                    ) from exc
                except Exception as exc:  # pylint: disable=broad-exception-caught
                    logger.error(f"XLSX extraction failed: {exc}")
                    raise RuntimeError(f"Spreadsheet extraction failed: {exc}") from exc

            if lower_path.endswith((".csv", ".tsv")):
                delimiter = "\t" if lower_path.endswith(".tsv") else ","
                rows = []
                with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
                    reader = csv.reader(f, delimiter=delimiter)
                    for row in reader:
                        values = [cell.strip() for cell in row if cell and cell.strip()]
                        if values:
                            rows.append("\t".join(values))

                if not rows:
                    return "Spreadsheet contains no extractable text."
                return "\n".join(rows)

            if lower_path.endswith(".xls"):
                raise RuntimeError(
                    "Legacy .xls is not supported. Please convert to .xlsx."
                )

            raise RuntimeError("Unsupported spreadsheet format.")
        finally:
            if self._is_ephemeral_download(sheet_path, local_path) and os.path.exists(
                local_path
            ):
                os.remove(local_path)


multimedia_service = MultimediaService()
